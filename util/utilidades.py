# -*- coding: utf-8 -*-
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import NoAlertPresentException
import os
import sys
from docx import Document
from docx.shared import Inches

def acessar(driver, url):
	driver.get(url)

def logar(driver, nomeScript, usuario, senha):
	driver.find_element_by_id("j_username").clear()
	driver.find_element_by_id("j_username").send_keys(usuario)
	driver.find_element_by_id("j_password").clear()
	campoSenha = driver.find_element_by_id("j_password")
	campoSenha.send_keys(senha)
	tirar_print(driver,nomeScript,"Entrando no sistema")
	campoSenha.send_keys(Keys.ENTER)
	time.sleep(2)

def sair(driver):
	driver.find_element_by_link_text("Sair").click()
	pass

def get_versao_sistema(driver):
	rodape = driver.find_element_by_id("rodape")
	texto = rodape.find_elements_by_tag_name("div")[0].find_elements_by_tag_name("p")[1].find_elements_by_tag_name("span")[0].text
	return texto

def esperar_elemento(driver, maneiraProcura, elemento, tempoLimite):
	WebDriverWait(driver, tempoLimite).until(
        EC.presence_of_element_located((maneiraProcura, elemento)))

def encontrar_elemento(driver, maneiraProcura, elemento):
	return driver.find_element(maneiraProcura,elemento)

def encontrar_elementos(driver, elemento, maneiraProcura):
	return driver.find_elements(maneiraProcura,elemento)

def clicar(driver, maneiraProcura, elemento):
	elemento = encontrar_elemento(driver, maneiraProcura, elemento)
	if (elemento != None):
		elemento.click()

def digitar(driver, maneiraProcura, elemento, texto):
	elemento = encontrar_elemento(driver,maneiraProcura,elemento)
	if (elemento != None):
		elemento.clear()
		elemento.send_keys(texto)

def get_diretorio_atual():
	return os.path.dirname(os.path.abspath(__file__))

def read(filename, encoding='utf-8', errors='strict'):
    with open(filename, 'rb') as f:
        return f.read().decode(encoding, errors=errors)

def get_parametros_script(nomeScript):
	nomeArquivoParametros = nomeScript.replace(".py","") + "_parametros.txt"
	diretorioAtual = get_diretorio_atual()
	conteudo = read(diretorioAtual + "\\" + nomeArquivoParametros)
	linhas = conteudo.split("\n")
	parametros = {}
	for linha in linhas:
		divisaoParametro = linha.split('=')
		if (len(divisaoParametro) > 1):
			parametros[divisaoParametro[0]] = divisaoParametro[1]

	return parametros

def inicializar_diretorio_prints(nomeScript):
	nomePasta = get_diretorio_pasta_prints(nomeScript)
	#Se a pasta nao existe ainda, cria
	if (not os.path.isdir(nomePasta)):
		os.mkdir(nomePasta)
	#Se ja existe, apaga os prints que ja existem
	else:
		prints = os.listdir(nomePasta)
		for file in prints:
			os.remove(nomePasta + '\\'+ file)

def get_diretorio_pasta_prints(nomeScript):
	diretorioAtual = get_diretorio_atual()
	return diretorioAtual + '\\Documentos\\'+ nomeScript.replace(".py","") + "_prints"

descricoes_prints = []

def tirar_print(driver,nomeScript,descricao):
	nomePasta = get_diretorio_pasta_prints(nomeScript) 
	numeroPrint = len(os.listdir(nomePasta)) + 1
	nomePrint = nomeScript.replace(".py","") + "_print_" + str(numeroPrint)
	#driver.get_screenshot_as_file(nomePasta +'/' + nomePrint + '.png')
	driver.save_screenshot(nomePasta +'/' + nomePrint + '.jpg')
	descricoes_prints.append(descricao)

def salvar_documento_interacao(nomeScript):
	nomePasta = get_diretorio_pasta_prints(nomeScript)
	documento = Document()
	nomeDoScript = nomeScript.replace(".py","")
	numeroPrints = len(os.listdir(nomePasta))
	for i in range(1,numeroPrints+1):
		documento.add_paragraph(descricoes_prints[i-1])
		documento.add_picture(get_diretorio_pasta_prints(nomeScript)+'\\'+nomeDoScript+"_print_"+str(i)+".jpg", width=Inches(6.25))

	documento.save(get_diretorio_pasta_prints(nomeScript)+'\\'+nomeDoScript+"_prints.docx")